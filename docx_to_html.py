#!/usr/bin/env python3
"""
docx_to_html.py — Konvertiert Word-Dokumente (.docx) in schöne HTML-Seiten.
LaTeX-Formeln (OMML oder $...$-Syntax) werden als PNG-Bilder eingebettet.

Abhängigkeiten installieren:
    pip install python-docx matplotlib pillow lxml

Verwendung:
    python docx_to_html.py dokument.docx
    python docx_to_html.py dokument.docx --output ausgabe.html
    python docx_to_html.py dokument.docx --dark
"""

import argparse
import base64
import io
import os
import re
import sys
from pathlib import Path
import colorsys

# ── Basis-Farbton: einzige Variable, alles wird daraus abgeleitet ─────────────
BASE_HUE = 180  # ← hier ändern um das gesamte Farbschema anzupassen (0–360)

def _hue_to_hex(h: int, s: float, l: float) -> str:
    """hsl(h 0-360, s 0-100, l 0-100) → #rrggbb  (identisch zu CSS hsl())"""
    r, g, b = colorsys.hls_to_rgb(h / 360, l / 100, s / 100)
    return f"#{int(r*255):02x}{int(g*255):02x}{int(b*255):02x}"


# ── Pflicht-Imports ────────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.oxml.ns import qn
    from lxml import etree
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.mathtext
except ImportError as e:
    print(f"[FEHLER] Fehlende Abhängigkeit: {e}")
    print("Bitte ausführen:  pip install python-docx matplotlib pillow lxml")
    sys.exit(1)


# ══════════════════════════════════════════════════════════════════════════════
#  LATEX → PNG (Base64)
# ══════════════════════════════════════════════════════════════════════════════

# Matplotlib unterstützt nur eine Teilmenge von LaTeX.
# Dieser Preprocessor konvertiert häufig verwendete, nicht unterstützte
# Befehle in äquivalente oder lesbare Formen.

_TEXT_RE   = re.compile(r"\\text\{([^}]*)\}")
_MATHRM_RE = re.compile(r"\\mathrm\{([^}]*)\}")
_MATHBF_RE = re.compile(r"\\mathbf\{([^}]*)\}")
_MATHIT_RE = re.compile(r"\\mathit\{([^}]*)\}")
_OPERATORNAME_RE = re.compile(r"\\operatorname\{([^}]*)\}")
_MBOX_RE   = re.compile(r"\\mbox\{([^}]*)\}")

# Befehl-Aliase: nicht unterstützt → Matplotlib-Äquivalent oder Unicode
_CMD_MAP = {
    r"\Rightarrow":   r"\Rightarrow",   # ist in matplotlib ok, aber als Fallback
    r"\rightarrow":   r"\rightarrow",
    r"\Leftarrow":    r"\Leftarrow",
    r"\leftarrow":    r"\leftarrow",
    r"\Leftrightarrow": r"\Leftrightarrow",
    r"\implies":      r"\Rightarrow",
    r"\iff":          r"\Leftrightarrow",
    r"\coloneqq":     r":=",
    r"\eqqcolon":     r"=:",
    r"\triangleq":    r"\triangleq",
    r"\because":      r"\because",
    r"\therefore":    r"\therefore",
    r"\varnothing":   r"\emptyset",
    r"\lvert":        r"|",
    r"\rvert":        r"|",
    r"\lVert":        r"\|",
    r"\rVert":        r"\|",
    r"\mathbb{R}":    r"\mathbb{R}",
    r"\mathbb{N}":    r"\mathbb{N}",
    r"\mathbb{Z}":    r"\mathbb{Z}",
    r"\mathbb{C}":    r"\mathbb{C}",
    r"\mathbb{Q}":    r"\mathbb{Q}",
    r"\hbar":         r"\hbar",
    r"\partial":      r"\partial",
    r"\nabla":        r"\nabla",
    r"\infty":        r"\infty",
    r"\dagger":       r"\dagger",
    r"\dag":          r"\dagger",
    r"\Box":          r"\Box",
    r"\square":       r"\square",
    # Umgebungen vereinfachen
    r"\begin{aligned}": r"",
    r"\end{aligned}":   r"",
    r"\begin{align}":   r"",
    r"\end{align}":     r"",
    r"\begin{equation}": r"",
    r"\end{equation}":  r"",
    r"\\":              r"\quad",   # Zeilenumbruch → Abstand
    r"\nonumber":       r"",
    r"\left":           r"\left",
    r"\right":          r"\right",
}


def _preprocess_latex(expr: str) -> str:
    """Bereitet einen LaTeX-Ausdruck für Matplotlib's mathtext vor.

    - Konvertiert \\text{...} → \\mathrm{...}  (Matplotlib versteht \\mathrm)
    - Wendet Befehl-Aliase an
    - Entfernt nicht unterstützte Umgebungen
    """
    # \\text{...} → \\mathrm{...}  (einzige verlässliche Textschrift in mathtext)
    expr = _TEXT_RE.sub(lambda m: r"\mathrm{" + m.group(1).replace(" ", r"\ ") + "}", expr)
    expr = _MBOX_RE.sub(lambda m: r"\mathrm{" + m.group(1).replace(" ", r"\ ") + "}", expr)
    # \\operatorname{...} → \\mathrm{...}
    expr = _OPERATORNAME_RE.sub(lambda m: r"\mathrm{" + m.group(1) + "}", expr)

    # Bekannte Aliase ersetzen (längste zuerst, um Teilstrings zu vermeiden)
    for old, new in sorted(_CMD_MAP.items(), key=lambda x: -len(x[0])):
        expr = expr.replace(old, new)

    return expr


def latex_to_png_base64(latex: str, fontsize: int = 14, dpi: int = 150,
                         dark: bool = False) -> str | None:
    """Rendert einen LaTeX-Ausdruck mit Matplotlib und gibt Base64-PNG zurück."""
    expr = latex.strip()
    # Math-Mode sicherstellen
    if not (expr.startswith("$") or expr.startswith(r"\[")):
        expr = f"${expr}$"

    # Preprocessor anwenden
    # Inhalt zwischen den $...$ extrahieren, transformieren, neu einwickeln
    inner = expr
    display = False
    if inner.startswith("$$") and inner.endswith("$$"):
        inner = inner[2:-2]
        display = True
    elif inner.startswith("$") and inner.endswith("$"):
        inner = inner[1:-1]
    elif inner.startswith(r"\[") and inner.endswith(r"\]"):
        inner = inner[2:-2]
        display = True

    inner = _preprocess_latex(inner)
    expr = f"$${inner}$$" if display else f"${inner}$"

    # Textfarbe aus BASE_HUE berechnet (selbe Werte wie CSS --text)
    fg = _hue_to_hex(BASE_HUE, 60, 88) if dark else _hue_to_hex(BASE_HUE, 50, 10)

    # Versuch 1: normales Rendering
    try:
        b64 = _render_matplotlib(expr, fontsize, dpi, fg)
        if b64:
            return b64
    except Exception:
        pass

    # Versuch 2: bei Fehler nochmal mit größerem Fontsize und vereinfacht
    try:
        # Alle verbleibenden unbekannten Befehle als \mathrm rendern
        safe_inner = re.sub(r"\\([A-Za-z]+)", lambda m: m.group(0), inner)
        safe_expr = f"${safe_inner}$"
        b64 = _render_matplotlib(safe_expr, fontsize, dpi, fg)
        if b64:
            return b64
    except Exception:
        pass

    # Versuch 3: als plain text rendern (kein Mathe-Modus)
    try:
        plain = re.sub(r"\\[A-Za-z]+\{([^}]*)\}", r"\1", inner)
        plain = re.sub(r"\\[A-Za-z]+", "", plain)
        plain = plain.replace("{", "").replace("}", "").replace("$", "").strip()
        if plain:
            b64 = _render_matplotlib(plain, fontsize, dpi, fg, math=False)
            if b64:
                print(f"  [INFO] Formel als Klartext gerendert: {plain!r}")
                return b64
    except Exception:
        pass

    print(f"  [WARN] Formel konnte nicht gerendert werden: {latex!r}")
    plt.close("all")
    return None


def _render_matplotlib(expr: str, fontsize: int, dpi: int, fg: str,
                        math: bool = True) -> str | None:
    """Internes Rendering via Matplotlib. Gibt Base64-PNG oder None zurück."""
    fig, ax = plt.subplots(figsize=(0.01, 0.01))
    fig.patch.set_alpha(0)
    ax.set_axis_off()
    t = ax.text(
        0.5, 0.5, expr,
        fontsize=fontsize,
        ha="center", va="center",
        color=fg,
        usetex=False,
        math_fontfamily="cm" if math else None,
    )
    fig.canvas.draw()
    bbox = t.get_window_extent(renderer=fig.canvas.get_renderer())
    pad = 8
    w = (bbox.width  + pad * 2) / dpi
    h = (bbox.height + pad * 2) / dpi
    fig.set_size_inches(max(w, 0.5), max(h, 0.3))
    ax.set_position([
        pad / (fig.get_figwidth()  * dpi),
        pad / (fig.get_figheight() * dpi),
        1 - 2 * pad / (fig.get_figwidth()  * dpi),
        1 - 2 * pad / (fig.get_figheight() * dpi),
    ])
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight",
                transparent=True, pad_inches=0.06)
    plt.close(fig)
    buf.seek(0)
    data = buf.read()
    if len(data) < 100:   # leeres Bild
        return None
    return base64.b64encode(data).decode("utf-8")


def formula_img_tag(latex: str, display: bool = False,
                    dark: bool = False) -> str:
    """Gibt ein <img>-Tag mit eingebettetem PNG zurück, oder Fallback-<code>."""
    b64 = latex_to_png_base64(latex, fontsize=15 if display else 13, dark=dark)
    cls = "formula-display" if display else "formula-inline"
    if b64:
        alt = latex.replace('"', "&quot;")
        return f'<img class="{cls}" src="data:image/png;base64,{b64}" alt="{alt}">'
    # Fallback: roher Code
    safe = latex.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return f'<code class="{cls}-fallback">{safe}</code>'


# ══════════════════════════════════════════════════════════════════════════════
#  OMML (Office Math Markup Language) → LaTeX-String
# ══════════════════════════════════════════════════════════════════════════════

OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

def omml_to_latex(elem) -> str:
    """Einfache OMML-zu-LaTeX-Konvertierung (häufigste Konstrukte)."""
    tag = etree.QName(elem).localname if elem.tag else ""

    def child_latex(e):
        return "".join(omml_to_latex(c) for c in e)

    if tag == "oMath":
        return child_latex(elem)
    if tag == "r":            # run → Text
        t = elem.find(f"{{{OMML_NS}}}t")
        return t.text or "" if t is not None else ""
    if tag == "f":            # Bruch
        num = elem.find(f"{{{OMML_NS}}}num")
        den = elem.find(f"{{{OMML_NS}}}den")
        n = child_latex(num) if num is not None else ""
        d = child_latex(den) if den is not None else ""
        return rf"\frac{{{n}}}{{{d}}}"
    if tag == "sSup":         # Hochstellen
        base = elem.find(f"{{{OMML_NS}}}e")
        sup  = elem.find(f"{{{OMML_NS}}}sup")
        b = child_latex(base) if base is not None else ""
        s = child_latex(sup)  if sup  is not None else ""
        return f"{b}^{{{s}}}"
    if tag == "sSub":         # Tiefstellen
        base = elem.find(f"{{{OMML_NS}}}e")
        sub  = elem.find(f"{{{OMML_NS}}}sub")
        b = child_latex(base) if base is not None else ""
        s = child_latex(sub)  if sub  is not None else ""
        return f"{b}_{{{s}}}"
    if tag == "sSubSup":
        base = elem.find(f"{{{OMML_NS}}}e")
        sub  = elem.find(f"{{{OMML_NS}}}sub")
        sup  = elem.find(f"{{{OMML_NS}}}sup")
        b = child_latex(base) if base is not None else ""
        lo = child_latex(sub) if sub  is not None else ""
        hi = child_latex(sup) if sup  is not None else ""
        return f"{b}_{{{lo}}}^{{{hi}}}"
    if tag == "rad":          # Wurzel
        deg = elem.find(f"{{{OMML_NS}}}deg")
        e   = elem.find(f"{{{OMML_NS}}}e")
        d = child_latex(deg) if deg is not None else ""
        base = child_latex(e) if e is not None else ""
        return rf"\sqrt[{d}]{{{base}}}" if d.strip() else rf"\sqrt{{{base}}}"
    if tag == "nary":         # Summe / Integral
        sub = elem.find(f"{{{OMML_NS}}}sub")
        sup = elem.find(f"{{{OMML_NS}}}sup")
        e   = elem.find(f"{{{OMML_NS}}}e")
        # Operator aus naryPr > chr
        naryPr = elem.find(f"{{{OMML_NS}}}naryPr")
        op = r"\sum"
        if naryPr is not None:
            chr_el = naryPr.find(f"{{{OMML_NS}}}chr")
            if chr_el is not None:
                c = chr_el.get(f"{{{OMML_NS}}}val", "")
                if "∫" in c or c == "\u222b": op = r"\int"
                elif "∏" in c or c == "\u220f": op = r"\prod"
        lo = f"_{{{child_latex(sub)}}}" if sub is not None else ""
        hi = f"^{{{child_latex(sup)}}}" if sup is not None else ""
        body = child_latex(e) if e is not None else ""
        return f"{op}{lo}{hi} {body}"
    if tag == "d":            # Klammern
        return rf"\left({child_latex(elem)}\right)"
    if tag in ("e", "num", "den", "sup", "sub", "deg", "lim", "oMathPara",
               "oMathParaPr", "naryPr", "fPr", "rPr", "sPr", "dPr"):
        return child_latex(elem)
    # Alles andere: Kinder verarbeiten
    return child_latex(elem)


# ══════════════════════════════════════════════════════════════════════════════
#  DOCX-Parsing
# ══════════════════════════════════════════════════════════════════════════════

INLINE_MATH_RE       = re.compile(r"\$(.+?)\$", re.DOTALL)
DISPLAY_MATH_RE      = re.compile(r"\$\$(.+?)\$\$", re.DOTALL)
# [ E = mc^2 ] oder [ \eta ] — eckige Klammern mit Mathe-Inhalt
# Matcht wenn Inhalt mind. eines von: \ ^ _ = + - * / < > { } | ~ oder Ziffer enthaelt
BRACKET_MATH_RE      = re.compile(
    r"\[\s*((?:[^\[\]]*(?:[\\^_=+\-*/<>{}|~]|(?<![,\s])\d)[^\[\]]*)+?)\s*\]"
)
# \[ ... \]  — Display-Math mit Backslash-Klammern
BACKSLASH_DISPLAY_RE = re.compile(r"\\\[(.+?)\\\]", re.DOTALL)


def escape_html(text: str) -> str:
    return (text
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;"))


def run_to_html(run) -> str:
    """Konvertiert einen docx-Run in HTML (Fett, Kursiv, Unterstrichen, …)."""
    text = run.text or ""
    if not text:
        return ""
    safe = escape_html(text)
    if run.bold:        safe = f"<strong>{safe}</strong>"
    if run.italic:      safe = f"<em>{safe}</em>"
    if run.underline:   safe = f"<u>{safe}</u>"
    # Schriftfarbe
    color = None
    if run.font.color and run.font.color.rgb:
        color = f"#{run.font.color.rgb}"
    if color:
        safe = f'<span style="color:{color}">{safe}</span>'
    return safe


def inline_math_to_html(text: str, dark: bool) -> str:
    """Ersetzt alle bekannten Formel-Notationen durch Formel-Images.

    Unterstuetzte Formate:
      $$...$$       -> Display-Formel
      $...$         -> Inline-Formel
      \\[...\\]    -> Display-Formel (LaTeX-Standard)
      [ \\cmd ]    -> Inline-Formel (Word-Format: [ \\eta ])
    """
    def repl_display(m):
        return formula_img_tag(m.group(1).strip(), display=True, dark=dark)

    def repl_inline(m):
        return formula_img_tag(m.group(1).strip(), display=False, dark=dark)

    text = DISPLAY_MATH_RE.sub(repl_display, text)       # $$...$$
    text = BACKSLASH_DISPLAY_RE.sub(repl_display, text)  # \[...\]
    text = INLINE_MATH_RE.sub(repl_inline, text)         # $...$
    text = BRACKET_MATH_RE.sub(repl_inline, text)        # [ \eta ]
    return text


def paragraph_to_html(para, dark: bool) -> str:
    """Wandelt einen Paragraph (inkl. OMML-Formeln) in einen HTML-Block."""
    style = para.style.name if para.style else ""

    # Heading-Level
    heading_match = re.match(r"Heading (\d)", style, re.IGNORECASE)
    level = int(heading_match.group(1)) if heading_match else None

    # Ausrichtung
    align_map = {"CENTER": "center", "RIGHT": "right", "JUSTIFY": "justify"}
    align = ""
    if para.alignment:
        align = align_map.get(str(para.alignment).split(".")[-1], "")

    # Alle Kinder iterieren: w:r (Runs) und m:oMath (Formeln) mischen
    parts = []
    for child in para._p:
        child_tag = etree.QName(child).localname

        # OMML-Formel (eingebettet in oMathPara oder direkt als oMath)
        if child.tag == f"{{{OMML_NS}}}oMathPara":
            for math_el in child.findall(f"{{{OMML_NS}}}oMath"):
                latex = omml_to_latex(math_el)
                parts.append(formula_img_tag(latex, display=True, dark=dark))
            continue
        if child.tag == f"{{{OMML_NS}}}oMath":
            latex = omml_to_latex(child)
            parts.append(formula_img_tag(latex, display=False, dark=dark))
            continue

        # Normaler Run (w:r)
        if child_tag == "r":
            # Text extrahieren
            t_els = child.findall(qn("w:t"))
            raw = "".join((t.text or "") for t in t_els)
            if not raw:
                continue
            # Formatierung aus w:rPr
            rpr = child.find(qn("w:rPr"))
            bold      = rpr is not None and rpr.find(qn("w:b"))   is not None
            italic    = rpr is not None and rpr.find(qn("w:i"))   is not None
            underline = rpr is not None and rpr.find(qn("w:u"))   is not None
            color_el  = rpr.find(qn("w:color")) if rpr is not None else None
            color     = None
            if color_el is not None:
                c = color_el.get(qn("w:val"), "")
                if c and c.lower() not in ("auto", "000000", ""):
                    color = f"#{c}"

            # Inline-Math ($…$) im Text ersetzen
            html = inline_math_to_html(escape_html(raw), dark)

            if bold:      html = f"<strong>{html}</strong>"
            if italic:    html = f"<em>{html}</em>"
            if underline: html = f"<u>{html}</u>"
            if color:     html = f'<span style="color:{color}">{html}</span>'
            parts.append(html)
            continue

        # Hyperlink (w:hyperlink)
        if child_tag == "hyperlink":
            link_text = "".join(
                (t.text or "")
                for r in child.findall(f".//{qn('w:r')}")
                for t in r.findall(qn("w:t"))
            )
            # Relationship-ID → URL (vereinfacht; URL bleibt "#" wenn nicht verfügbar)
            parts.append(f'<a href="#">{escape_html(link_text)}</a>')

    inner = "".join(parts).strip()
    if not inner:
        return '<p class="empty-para">&nbsp;</p>'

    style_attr = f' style="text-align:{align}"' if align else ""

    if level:
        return f"<h{level}{style_attr}>{inner}</h{level}>"

    # Listenabsatz
    if "List" in style:
        list_cls = "ol-item" if "Number" in style else "ul-item"
        return f'<li class="{list_cls}">{inner}</li>'

    # Blockzitat
    if "Quote" in style or "Block" in style:
        return f"<blockquote>{inner}</blockquote>"

    return f"<p{style_attr}>{inner}</p>"


def table_to_html(table) -> str:
    rows_html = []
    for i, row in enumerate(table.rows):
        cells_html = []
        for cell in row.cells:
            cell_content = "\n".join(
                paragraph_to_html(p, dark=False)
                for p in cell.paragraphs if p.text.strip()
            )
            tag = "th" if i == 0 else "td"
            cells_html.append(f"<{tag}>{cell_content or '&nbsp;'}</{tag}>")
        rows_html.append(f"<tr>{''.join(cells_html)}</tr>")
    return f"<table>{''.join(rows_html)}</table>"


# ══════════════════════════════════════════════════════════════════════════════
#  Bild-Extraktion
# ══════════════════════════════════════════════════════════════════════════════

def image_inline_b64(rel_id: str, doc) -> str | None:
    """Gibt ein eingebettetes Bild als Base64-Data-URL zurück."""
    try:
        rel = doc.part.rels.get(rel_id)
        if rel and "image" in rel.reltype:
            img_bytes = rel.target_part.blob
            mt = rel.target_part.content_type or "image/png"
            b64 = base64.b64encode(img_bytes).decode()
            return f"data:{mt};base64,{b64}"
    except Exception:
        pass
    return None


# ══════════════════════════════════════════════════════════════════════════════
#  HTML-Template
# ══════════════════════════════════════════════════════════════════════════════

CSS_LIGHT = """
:root {
  /* Einzige Basis-Variable — alle Farben automatisch abgeleitet */
  --hue:     {hue};
  --hue2:    calc(var(--hue) + 150);  /* Komplementärton */

  --bg:        hsl(var(--hue), 20%,  96%);
  --surface:   hsl(var(--hue), 15%,  99%);
  --text:      hsl(var(--hue), 50%,  10%);
  --muted:     hsl(var(--hue), 30%,  42%);
  --accent:    hsl(var(--hue), 80%,  32%);
  --accent2:   hsl(var(--hue2), 75%, 40%);
  --border:    hsl(var(--hue), 30%,  82%);
  --code-bg:   hsl(var(--hue), 25%,  91%);
  --shadow:    hsl(var(--hue), 60%,  10%, 0.10);
}
"""

CSS_DARK = """
:root {
  /* Einzige Basis-Variable — alle Farben automatisch abgeleitet */
  --hue:     {hue};
  --hue2:    calc(var(--hue) + 150);  /* Komplementärton */

  --bg:        hsl(var(--hue), 60%,   7%);
  --surface:   hsl(var(--hue), 55%,  12%);
  --text:      hsl(var(--hue), 60%,  88%);
  --muted:     hsl(var(--hue), 40%,  60%);
  --accent:    hsl(var(--hue), 90%,  60%);
  --accent2:   hsl(var(--hue2), 80%, 65%);
  --border:    hsl(var(--hue), 50%,  20%);
  --code-bg:   hsl(var(--hue), 55%,  16%);
  --shadow:    hsl(var(--hue), 70%,   3%, 0.55);
}
"""

CSS_BASE = """
*, *::before, *::after { box-sizing: border-box; margin: 0; padding: 0; }

body {
  font-family: 'Lora', 'Georgia', serif;
  background: var(--bg);
  color: var(--text);
  line-height: 1.85;
  font-size: 17px;
  padding: 0 1rem 4rem;
}

.page-wrapper {
  max-width: 780px;
  margin: 0 auto;
  background: var(--surface);
  box-shadow: 0 4px 40px var(--shadow);
  border-radius: 12px;
  overflow: hidden;
}

.doc-header {
  background: linear-gradient(135deg, var(--accent) 0%, var(--accent2) 100%);
  padding: 3rem 3rem 2.5rem;
  color: var(--surface);
}
.doc-header h1 {
  font-family: 'Playfair Display', serif;
  font-size: 2.2rem;
  font-weight: 700;
  margin: 0;
  line-height: 1.2;
  color: var(--surface);
}
.doc-meta {
  margin-top: 0.6rem;
  font-size: 0.85rem;
  opacity: 0.75;
  font-family: 'JetBrains Mono', monospace;
}

.content {
  padding: 2.5rem 3rem 3rem;
}

/* ── Headings ── */
h1, h2, h3, h4, h5, h6 {
  font-family: 'Playfair Display', serif;
  color: var(--text);
  margin-top: 2rem;
  margin-bottom: 0.6rem;
  line-height: 1.3;
}
h1 { font-size: 2rem; border-bottom: 3px solid var(--accent); padding-bottom: 0.3rem; }
h2 { font-size: 1.55rem; color: var(--accent); }
h3 { font-size: 1.25rem; }
h4 { font-size: 1.05rem; font-style: italic; }

/* ── Paragraphs ── */
p {
  margin-bottom: 1rem;
  color: var(--text);
}
p.empty-para { margin-bottom: 0.4rem; }

/* ── Inline formatting ── */
strong { font-weight: 700; color: var(--accent); }
em     { font-style: italic; }
u      { text-decoration: underline; text-decoration-color: var(--accent2); }
a      { color: var(--accent); text-decoration: none; border-bottom: 1px solid currentColor; }
a:hover{ opacity: 0.75; }

/* ── Formeln ── */
.formula-display {
  display: block;
  margin: 1.6rem auto;
  max-width: 100%;
  filter: drop-shadow(0 2px 4px var(--shadow));
}
.formula-inline {
  display: inline;
  vertical-align: middle;
  margin: 0 2px;
}
.formula-display-fallback,
.formula-inline-fallback {
  font-family: 'JetBrains Mono', monospace;
  background: var(--code-bg);
  border-radius: 4px;
  padding: 0.1em 0.4em;
  font-size: 0.88em;
  color: var(--accent2);
  display: inline-block;
  margin: 0 2px;
}

/* ── Blockquote ── */
blockquote {
  border-left: 4px solid var(--accent);
  padding: 0.8rem 1.2rem;
  margin: 1.2rem 0;
  background: var(--code-bg);
  border-radius: 0 8px 8px 0;
  font-style: italic;
  color: var(--muted);
}

/* ── Listen ── */
li { margin-left: 1.8rem; margin-bottom: 0.35rem; }
li.ul-item { list-style-type: disc; }
li.ol-item { list-style-type: decimal; }

/* ── Tabellen ── */
table {
  width: 100%;
  border-collapse: collapse;
  margin: 1.5rem 0;
  font-size: 0.95rem;
  border-radius: 8px;
  overflow: hidden;
  box-shadow: 0 2px 12px var(--shadow);
}
th {
  background: var(--accent);
  color: var(--surface);
  font-family: 'Playfair Display', serif;
  font-weight: 700;
  padding: 0.75rem 1rem;
  text-align: left;
}
td {
  padding: 0.6rem 1rem;
  border-bottom: 1px solid var(--border);
  vertical-align: top;
}
tr:last-child td { border-bottom: none; }
tr:nth-child(even) td { background: var(--code-bg); }

/* ── Bilder ── */
img.doc-image {
  max-width: 100%;
  border-radius: 8px;
  box-shadow: 0 4px 20px var(--shadow);
  margin: 1.2rem 0;
  display: block;
}

/* ── Footer ── */
.doc-footer {
  text-align: center;
  padding: 1.2rem;
  font-size: 0.78rem;
  color: var(--muted);
  border-top: 1px solid var(--border);
  font-family: 'JetBrains Mono', monospace;
}

/* ── Responsive ── */
@media (max-width: 600px) {
  .content { padding: 1.5rem 1.2rem 2rem; }
  .doc-header { padding: 2rem 1.2rem 1.5rem; }
  h1 { font-size: 1.6rem; }
}
"""

HTML_TEMPLATE = """\
<!DOCTYPE html>
<html lang="de">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{title}</title>
  <link rel="preconnect" href="https://fonts.googleapis.com">
  <link href="https://fonts.googleapis.com/css2?family=Lora:ital,wght@0,400;0,600;1,400&family=Playfair+Display:wght@700;800&family=JetBrains+Mono:wght@400&display=swap" rel="stylesheet">
  <style>
{theme_css}
{base_css}
  </style>
</head>
<body>
  <div class="page-wrapper">
    <header class="doc-header">
      <h1>{title}</h1>
      <div class="doc-meta">Konvertiert von {source} · {timestamp}</div>
    </header>
    <div class="content">
{body}
    </div>
    <footer class="doc-footer">Generiert mit docx_to_html.py  © Zami & Jo - 2026</footer>
  </div>
</body>
</html>
"""


# ══════════════════════════════════════════════════════════════════════════════
#  Haupt-Konvertierung
# ══════════════════════════════════════════════════════════════════════════════

def convert(docx_path: str, output_path: str | None = None, dark: bool = False) -> str:
    from datetime import datetime

    print(f"📄  Lade {docx_path} …")
    doc = Document(docx_path)

    title = Path(docx_path).stem
    # Versuche, Titel aus erstem Heading zu lesen
    for para in doc.paragraphs:
        if para.style and "Heading 1" in para.style.name and para.text.strip():
            title = para.text.strip()
            break

    print("🔄  Konvertiere Inhalt …")
    body_parts = []
    in_ul = False
    in_ol = False

    for block in doc.element.body:
        tag = etree.QName(block).localname

        # ── Absatz ──────────────────────────────────────────────────────────
        if tag == "p":
            from docx.text.paragraph import Paragraph as DocxPara
            para = DocxPara(block, doc)
            html = paragraph_to_html(para, dark=dark)

            is_li = html.startswith("<li")
            is_ol = 'class="ol-item"' in html

            if is_li:
                if not (in_ul or in_ol):
                    tag_open = "<ol>" if is_ol else "<ul>"
                    body_parts.append(tag_open)
                    in_ul, in_ol = not is_ol, is_ol
                body_parts.append(html)
            else:
                if in_ul:  body_parts.append("</ul>"); in_ul = False
                if in_ol:  body_parts.append("</ol>"); in_ol = False
                body_parts.append(html)

        # ── Tabelle ─────────────────────────────────────────────────────────
        elif tag == "tbl":
            if in_ul: body_parts.append("</ul>"); in_ul = False
            if in_ol: body_parts.append("</ol>"); in_ol = False
            from docx.table import Table as DocxTable
            tbl = DocxTable(block, doc)
            body_parts.append(table_to_html(tbl))

    if in_ul: body_parts.append("</ul>")
    if in_ol: body_parts.append("</ol>")

    body_html = "\n      ".join(body_parts)
    import re as _re
    _css_tpl = CSS_DARK if dark else CSS_LIGHT
    theme_css  = _re.sub(r"\{hue\}", str(BASE_HUE), _css_tpl)

    html = HTML_TEMPLATE.format(
        title=escape_html(title),
        source=escape_html(Path(docx_path).name),
        timestamp=datetime.now().strftime("%d.%m.%Y %H:%M"),
        theme_css=theme_css,
        base_css=CSS_BASE,
        body=body_html,
    )

    if output_path is None:
        output_path = Path(docx_path).with_suffix(".html")

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html)

    print(f"✅  HTML gespeichert: {output_path}")
    return str(output_path)


# ══════════════════════════════════════════════════════════════════════════════
#  CLI
# ══════════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="Konvertiert ein Word-Dokument (.docx) in eine schöne HTML-Seite."
    )
    parser.add_argument("docx", help="Pfad zur .docx-Datei")
    parser.add_argument("-o", "--output", help="Ausgabe-HTML-Datei (optional)", default=None)
    parser.add_argument("--dark", action="store_true", help="Dunkles Theme verwenden")
    args = parser.parse_args()

    if not os.path.isfile(args.docx):
        print(f"[FEHLER] Datei nicht gefunden: {args.docx}")
        sys.exit(1)

    convert(args.docx, args.output, dark=args.dark)


if __name__ == "__main__":
    main()
